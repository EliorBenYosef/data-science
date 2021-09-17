import re
import PyPDF2
import docx
from nltk.stem.porter import PorterStemmer
from nltk.stem.wordnet import WordNetLemmatizer
# from unidecode import unidecode
from sklearn.metrics.pairwise import cosine_similarity, cosine_distances
from nltk.tag import pos_tag
from nltk.tokenize import word_tokenize
import numpy as np
import os


# from nltk import download
# download('stopwords')
# download('wordnet')
# download('averaged_perceptron_tagger')  # for pos_tag
# download('punkt')  # for word_tokenize

from nltk.corpus import stopwords
from nltk.corpus import wordnet


class TextCleaner:

    # re.sub(r"\s+", "", text)  # remove all spaces
    # re.sub(r"^\s+", "", text)  # remove left (leading) space
    # re.sub(r"\s+$", "", text)  # remove right (trailing) space
    # re.sub(r"^\s+|\s+$", "", text)  # remove both
    # re.sub(' +', ' ',text)  # replace multiple spaces with a single space

    def __init__(self, remove_stopwords=True, apply_normalization=True, lemmatize=True):
        """
        stopwords - irrelevant words to exclude when cleaning the texts:
            articles (the, this, a, an, ...)
            conjunctions (and, but, or, while, because, ...)
            pronouns (I, you, ...)
            prepositions (by, with, about, until, ...)
            'not'
            ...
        http://www.butte.edu/departments/cas/tipsheets/grammar/parts_of_speech.html

        stemming - taking only the root of the word.
        """
        # self.text =
        # self.text_clean =

        self.remove_stopwords = remove_stopwords
        if self.remove_stopwords:
            self.stop_words = stopwords.words('english')  # will be removed from the text
            # words = []
            # for word in self.stop_words:
            #     if "\'" in word:
            #         words.append(word)
            # print(words)

        self.apply_normalization = apply_normalization
        self.lemmatize = lemmatize
        if self.apply_normalization:
            self.ps = PorterStemmer()
            self.wnl = WordNetLemmatizer()
            # Map POS (Part-of-Speech) tag to pos arg lemmatize(pos='') accepts:
            self.pos_tag_dict = {'NN': wordnet.NOUN,
                                 'VB': wordnet.VERB,
                                 'JJ': wordnet.ADJ,
                                 'RB': wordnet.ADV}

    def print_tokens(self, text):
        print(pos_tag(word_tokenize(text)))

    def lemmatize_text(self, text, print_tokens=False):
        """
        ADJ, ADJ_SAT, ADV, NOUN, VERB = "a", "s", "r", "n", "v"
        https://www.learntek.org/blog/categorizing-pos-tagging-nltk-python/

        reminder: yield returns a generator, can be used in list comprehension:
            [lem_word for lem_word in lemmatize_text(text)]
        """
        # text = text.replace("/", ' or ')
        # text = text.replace("\\", ' or ')
        # # text = text.replace("'s", '')
        # # text = text.replace("’s", '')

        if print_tokens:
            print(pos_tag(word_tokenize(text)))

        # text = "We’re looking for an exceptional Deep Learning (DL) Engineer"  # TODO: remove
        for word, tag in pos_tag(word_tokenize(text)):
            if tag.startswith('NN'):  # NOUN
                # NN noun, singular ‘desk’, ’dog’
                # NNS noun plural ‘desks’, ‘dogs’
                # NNP proper noun, singular ‘Harrison’
                # NNPS proper noun, plural ‘Americans’
                yield self.wnl.lemmatize(word, pos='n')  # wordnet.NOUN
            elif tag.startswith('VB'):  # VERB
                # VB verb, base form take
                # VBD verb, past tense took
                # VBG verb, gerund/present participle taking
                # VBN verb, past participle taken
                # VBP verb, sing. present, non-3d take
                # VBZ verb, 3rd person sing. present takes
                yield self.wnl.lemmatize(word, pos='v')  # wordnet.VERB
            elif tag.startswith('JJ'):  # ADJ
                # JJ adjective ‘big’, ’good’
                # JJR adjective, comparative ‘bigger’, ‘better’
                # JJS adjective, superlative ‘biggest’
                yield self.wnl.lemmatize(word, pos='a')  # wordnet.ADJ
            elif tag.startswith('RB'):  # ADV
                # RB adverb very, silently,
                # RBR adverb, comparative better
                # RBS adverb, superlative best
                yield self.wnl.lemmatize(word, pos='r')  # wordnet.ADV
            else:
                yield word

    def clean_lemmatized_text(self, text):
        text = text.replace(" ’ ", "'")
        text = text.replace(" ' ", "'")

        text = text.lower()
        text = self.remove_stopwords_f(text)
        text = re.sub('[^A-Za-z]+', ' ', text)
        return text

    def get_clean_lemmatized_text(self, text):
        text = ' '.join(self.lemmatize_text(text))
        text = self.clean_lemmatized_text(text)
        return text

    def extract_keywords(self, text, print_tokens=False):
        # https://www.learntek.org/blog/categorizing-pos-tagging-nltk-python/

        text = text.replace("/", ' or ')
        text = text.replace("\\", ' or ')
        text = text.replace("'s", '')
        text = text.replace("’s", '')

        if print_tokens:
            print(pos_tag(word_tokenize(text)))

        keywords = set()
        phrase = []
        jj = ''
        for word, tag in pos_tag(word_tokenize(text)):
            if tag == 'NNP':  # keep 'NNP' streaks
                phrase.append(word.lower())
                jj = ''
            elif tag == 'JJ':
                # remember last 'JJ' for 'JJ'+'NN' combo (keep both together) \ 'JJ'+'NNS' combo (keep both apart)
                jj = word.lower()
            elif tag == 'NN':
                if jj:
                    # 'JJ'+'NN' combo - keep both together
                    # phrase.append(jj + ' ' + self.wnl.lemmatize(word.lower(), pos='n'))
                    phrase.append(jj + ' ' + word.lower())
                    keywords.add(' '.join(phrase))
                    phrase = []
                    jj = ''
                else:
                    if phrase:
                        keywords.add(' '.join(phrase))
                        phrase = []
                    keywords.add(self.wnl.lemmatize(word.lower(), pos='n'))  # algorithms --> algorithm
            elif tag == 'NNS' and jj:
                # 'JJ'+'NNS' combo - keep both apart
                keywords.add(jj)
                # jj - conversational, exceptional, minimum\preferred, statistical, new
                keywords.add(self.wnl.lemmatize(word.lower(), pos='n'))
                # word - technologies, skills, qualifications, models\techniques, technologies
                phrase = []
                jj = ''
            else:
                # ignores: 'PRP', 'PRP$', 'VBP', 'IN', 'DT', 'TO', 'VB', 'CC', 'MD', 'VBG', 'VBZ', 'WDT', 'VBN',
                #   'RB', 'WP', '(', ')', ',', '.', ':'
                if phrase:
                    keywords.add(' '.join(phrase))
                    phrase = []
                jj = ''

        # # ['conversational technology', 'exceptional skill', 'new technology', 'state-of-the-art machine',
        #    'statistical model', 'statistical technique']
        # removed_keywords = ['do', 'world', 'advantage', 'working',
        #                     'exceptional skill',
        #                     'minimum qualification', 'preferred qualification',
        #                     'minimum requirement', 'preferred requirement']
        # ['conversational', 'exceptional', 'statistical']
        removed_keywords = ['do', 'world', 'advantage', 'working',
                            'model', 'new', 'skill', 'technique', 'technology',
                            'minimum', 'preferred', 'qualification', 'requirement']
        for element in removed_keywords:
            if element in keywords:
                keywords.remove(element)

        return sorted(keywords)

    def remove_stopwords_f(self, text, keep_neg_words=False):
        text = text.replace("’", "'")

        if keep_neg_words:
            # break negative contractions:  # ain't ??
            text = text.replace("can't", 'can not')
            text = text.replace("shan't", 'shall not')
            text = text.replace("won't", 'will not')
            text = text.replace("n't", ' not')

        text = text.replace("i'm", '')  # 'i am'
        text = text.replace("'s", '')  # 'he is', 'he has'  # included in stop_eords: "she's", "it's"
        text = text.replace("'re", '')  # ' are'  # included in stop_eords: "you're"
        text = text.replace("'ve", '')  # ' have'  # included in stop_eords: "you've", "should've"
        text = text.replace("'d", '')  # ' had', ' would'  # included in stop_eords: "you'd"
        text = text.replace("'ll", '')  # ' will'  # included in stop_eords: "you'll", "that'll"
        text = text.replace(" can't ", ' ')
        text = text.replace(" shall ", ' ')
        text = text.replace(" us ", ' ')

        text = [word for word in text.split() if word not in set(self.stop_words)]
        text = ' '.join(text)  # convert list back to string
        return text

    def get_wordnet_pos(self, word):
        """
        get a single word's wordnet POS (Part-of-Speech) tag.
        """
        # token = word_tokenize(word)
        base_tag = pos_tag([word])[0][1][:2]
        return self.pos_tag_dict.get(base_tag, wordnet.NOUN)

    def normalize_text(self, text):
        """
        Goal:
        - Identifying a canonical representative (or root) for a set of related word forms.
        - Reducing inflectional forms and sometimes derivationally related forms of a word to a common base form.

        Special cases of text normalization:
        * Stemming - a crude heuristic process that chops off the ends of words in the hope of achieving the goal
            correctly most of the time, and often includes the removal of derivational affixes.
            strips everything to the bare root. removes:
            any addition to verbs: 'e', 'es', 'ed', 'ing', ...
            plural 's', technology 'y', 'ies'
            'ational' (conversational), 'ation' (expectation), 'ional' (exceptional), 'ication' (qualification)
            'er' (engineer)
        * Lemmatization - doing things properly with the use of a vocabulary and morphological analysis of words,
            normally aiming to remove inflectional endings only and to return the base or dictionary form of a word,
            which is known as the lemma.
            removes plural form
        """
        if self.lemmatize:  # Lemmatization (note: this lemmatizes each word separately, and not as part of a sentence)
            # lemmatizes all words as nouns
            text = [self.wnl.lemmatize(word) for word in text.split()]
            # # lemmatizes all words as their (disconnected) type
            # text = [self.wnl.lemmatize(word, pos=self.get_wordnet_pos(word)) for word in text.split()]
        else:  # Stemming
            text = [self.ps.stem(word) for word in text.split()]

        text = ' '.join(text)  # convert list back to string

        return text

    def clean_review(self, text):
        """
        Cleans a single review (simplifies it as much as possible)
        """
        text = text.lower()  # lowercase capital letters

        if self.remove_stopwords:
            text = self.remove_stopwords_f(text, keep_neg_words=True)

        text = re.sub('[^a-zA-Z ]', repl=' ', string=text)  # remove non-letters characters (i.e. punctuation)

        text = re.sub(' +', ' ', text)  # remove extra whitespace

        if self.apply_normalization:
            text = self.normalize_text(text)

        return text

    def clean_resume(self, text):
        """
        Cleans a single resume (resume text)
        """
        text = text.lower()  # lowercase capital letters

        text = re.sub('httpS+s*', '', text)  # remove URLs
        text = re.sub('RT|cc', '', text)  # remove RT and cc
        text = re.sub('#S+', '', text)  # remove hashtags
        text = re.sub('@S+', '', text)  # remove mentions

        if self.remove_stopwords:
            text = self.remove_stopwords_f(text)

        text = re.sub('[^a-zA-Z ]', repl=' ', string=text)  # remove non-letters characters (i.e. punctuation)

        # text = re.sub('[%s]' % re.escape("""!"#$%&'()*+,-./:;<=>?@[]^_`{|}~"""), '', text)  # remove punctuation
        # text = re.sub(r'[^\x00-\x7f]', '', text)  # remove non-ASCII characters
        # # # Replace non-ASCII characters with their most alike representation (doesn't work):
        # # text = unidecode(unicode(text, encoding="utf-8"))

        text = re.sub(' +', ' ', text)  # remove extra whitespace

        if self.apply_normalization:
            text = self.normalize_text(text)

        return text

    def clean_job_description(self, text):
        text = text.lower()  # lowercase capital letters

        if self.remove_stopwords:
            text = self.remove_stopwords_f(text)

        # text = re.sub('\W+', ' ', text)  # Select only alphanumeric characters (letters & numbers)
        text = re.sub('[^A-Za-z]+', ' ', text)  # select only alphabet characters (letters only)

        if self.apply_normalization:
            text = self.normalize_text(text)

        return text

    @staticmethod
    def get_candidate_name_from_filename(file_path):
        file_name = os.path.splitext(os.path.basename(file_path))[0]
        text = file_name

        # clean_resume_file_name:
        text = text.lower()
        text = text.replace('resume', '').replace('cv', '')
        text = text.replace('english', '').replace('eng', '')
        text = re.sub('[^a-zA-Z ]', repl=' ', string=text)  # remove non-letters characters (i.e. punctuation)
        text = re.sub(' +', ' ', text)  # remove extra whitespace
        text = re.sub(r"^\s+|\s+$", "", text)  # remove both

        candidate_name = text.title()
        return candidate_name


class SimilarityMetrics:
    """
    Distance functions

    Bregman Divergence
    - Euclidean distance
        - Squared Euclidean distance
    - Squared Mahalanobis distance
    - Kullback–Leibler divergence (KL divergence)
    - Itakura–Saito distance
    """

    def __init__(self, v_base, v_compare: list):
        self.v_base = v_base
        self.v_compare = v_compare

    def euc_dist(self, squared=True):
        """
        Euclidean distance
        - Squared Euclidean distance - more frequently used
        """

    def cos_dist(self):
        """
        Cosine distance
        """

    def cos_sim(self):
        """
        Cosine similarity
        """
        similarity_list = []
        for v in self.v_compare:
            similarity_list.append(cosine_similarity(self.v_base, np.expand_dims(v, axis=0))[0][0])
        return similarity_list

    def l1_dist(self):
        """
        L1 distance
        """

    def l2_dist(self):
        """
        L2 distance
        """

    def kl_div(self, prob_base, prob_compare: list):
        """
        KL divergence
        """

    def js_div(self, prob_base, prob_compare: list):
        """
        JS divergence
        """

    def was_dist(self, prob_base, prob_compare: list):
        """
        Wasserstein distance
        """


class TextExtractor:

    @staticmethod
    def read_pdf(file_path):
        file_reader = PyPDF2.PdfFileReader(open(file_path, 'rb'))  # rb - read binary
        text = []
        for i in range(file_reader.getNumPages()):
            text.append(file_reader.getPage(i).extractText())
        text = ' '.join(text)  # convert list back to string
        text = text.replace('\n', '')
        return text

    @staticmethod
    def read_docx(file_path):
        doc = docx.Document(file_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        text = '\n'.join(text)  # convert list back to string
        text = text.replace('\n', ' . ')

        # # for tables:
        # tables_text = []
        # for table in doc.tables:
        #     for row in table.rows:
        #         rowText = ''
        #         for cell in row.cells:
        #             rowText = rowText + ' ' + cell.text
        #         tables_text.append(tables_text + ' ' + rowText)

        return text

    # @staticmethod
    # def read_docx_via_textract(file_path):
    #     text = textract.process(file_path)
    #     lower_case_string = str(text.decode('utf-8')).lower()
    #     # final_string = re.sub('[^a-zA-Z0-9 \n]', '', lower_case_string)
    #     return lower_case_string


if __name__ == '__main__':
    tc = TextCleaner()

    jd_text = TextExtractor.read_docx('../../datasets/per_field/nlp/job_description.docx')
    jd_text_lem_w = tc.clean_job_description(jd_text)
    jd_text_lem_t = tc.get_clean_lemmatized_text(jd_text)
    jd_keywords = tc.extract_keywords(jd_text)
